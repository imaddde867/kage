"""Local artifact connectors: safe read-only file discovery and extraction.

This module adds first-class local artifact tooling so the agent can inspect
user-requested files (for example PDF specs under Downloads) without relying on
shell fallbacks.

Safety model:
- Read-only operations only
- Paths must stay under configured safe roots
- Sensitive roots and filenames are blocked by default
- No writes, deletes, or shell execution
"""
from __future__ import annotations

import csv
import io
import os
from pathlib import Path
import re
from typing import Any

import config as _config
from core.agent.tool_base import Tool, ToolOutcome, ToolResult

try:
    from pypdf import PdfReader as _PdfReader  # type: ignore[import]
except ImportError:
    _PdfReader = None  # type: ignore[assignment]

try:
    from docx import Document as _DocxDocument  # type: ignore[import]
except ImportError:
    _DocxDocument = None  # type: ignore[assignment]

try:
    from openpyxl import load_workbook as _load_workbook  # type: ignore[import]
except ImportError:
    _load_workbook = None  # type: ignore[assignment]

_DEFAULT_MAX_CHARS = 8000
_MAX_ALLOWED_CHARS = 20000
_DEFAULT_FIND_MAX_RESULTS = 8
_MAX_FIND_RESULTS = 20
_DEFAULT_FIND_MAX_FILES = 50000
_DEFAULT_PDF_MAX_PAGES = 30
_MAX_PDF_PAGES = 120
_DEFAULT_SHEET_MAX_ROWS = 120
_MAX_SHEET_ROWS = 500
_DEFAULT_SHEET_MAX_SHEETS = 3
_MAX_SHEET_SHEETS = 8
_DEFAULT_MAX_FILE_BYTES = 8 * 1024 * 1024

_PDF_EXTENSIONS = {".pdf"}
_DOCX_EXTENSIONS = {".docx"}
_SHEET_EXTENSIONS = {".csv", ".tsv", ".xlsx", ".xlsm", ".xltx", ".xltm"}

_DENIED_FILENAMES = {
    ".env",
    ".env.local",
    ".env.production",
    ".env.development",
    "id_rsa",
    "id_dsa",
    "id_ecdsa",
    "id_ed25519",
    "authorized_keys",
    "known_hosts",
}

_WS_RE = re.compile(r"\s+")
_TOKEN_RE = re.compile(r"[A-Za-z0-9]+")


def _clamp_int(value: Any, default: int, *, low: int, high: int) -> int:
    try:
        parsed = int(value)
    except Exception:
        return default
    return max(low, min(parsed, high))


def _tool_result(
    *,
    tool_name: str,
    content: str,
    is_error: bool = False,
    status: str | None = None,
    structured: dict[str, Any] | None = None,
    sources: list[str] | None = None,
    retryable: bool | None = None,
) -> ToolResult:
    if retryable is None:
        retryable = is_error
    if status is None:
        status = "error" if is_error else "ok"
    return ToolResult(
        tool_name=tool_name,
        content=content,
        is_error=is_error,
        outcome=ToolOutcome(
            status=status,
            structured=structured,
            sources=list(sources or []),
            retryable=retryable,
            side_effects=False,
        ),
    )


def _normalize_space(text: str) -> str:
    return _WS_RE.sub(" ", text).strip()


def _candidate_path(raw: str) -> Path | None:
    value = str(raw or "").strip()
    if not value:
        return None
    candidate = Path(value).expanduser()
    if candidate.is_absolute():
        return candidate.resolve()
    return (Path.cwd() / candidate).resolve()


def _safe_roots(settings: Any) -> tuple[Path, ...]:
    raw = tuple(getattr(settings, "local_artifact_safe_roots", ()) or ())
    if not raw:
        raw = (str(Path.cwd()), str(Path.home()))
    roots: list[Path] = []
    seen: set[str] = set()
    for item in raw:
        path = _candidate_path(str(item))
        if not path:
            continue
        key = str(path)
        if key in seen:
            continue
        seen.add(key)
        roots.append(path)
    return tuple(roots)


def _deny_roots(settings: Any) -> tuple[Path, ...]:
    raw = tuple(getattr(settings, "local_artifact_deny_roots", ()) or ())
    if not raw:
        raw = (
            "~/.ssh",
            "~/.gnupg",
            "~/.aws",
            "~/.kube",
            "~/Library/Keychains",
            "~/Library/Application Support",
            "/etc",
            "/private",
            "/System",
            "/usr",
            "/bin",
            "/sbin",
        )
    roots: list[Path] = []
    seen: set[str] = set()
    for item in raw:
        path = _candidate_path(str(item))
        if not path:
            continue
        key = str(path)
        if key in seen:
            continue
        seen.add(key)
        roots.append(path)
    return tuple(roots)


def _is_within(path: Path, root: Path) -> bool:
    return path == root or root in path.parents


def _path_access_reason(path: Path, *, safe_roots: tuple[Path, ...], deny_roots: tuple[Path, ...]) -> str | None:
    if path.name in _DENIED_FILENAMES:
        return f"Blocked sensitive filename: {path.name}"

    in_safe_root = any(_is_within(path, root) for root in safe_roots)
    if not in_safe_root:
        allowed = ", ".join(str(root) for root in safe_roots)
        return f"Path is outside safe roots. Allowed roots: {allowed}"

    denied_matches = [denied for denied in deny_roots if _is_within(path, denied)]
    if denied_matches:
        # Explicit safe roots can override deny roots only for their own subtree.
        safe_override = any(
            _is_within(root, denied) and _is_within(path, root)
            for denied in denied_matches
            for root in safe_roots
        )
        if not safe_override:
            return f"Blocked sensitive path by policy: {path}"

    return None


def _allow_path(path: Path, *, safe_roots: tuple[Path, ...], deny_roots: tuple[Path, ...]) -> tuple[bool, str]:
    reason = _path_access_reason(path, safe_roots=safe_roots, deny_roots=deny_roots)
    if reason:
        return False, reason
    return True, ""


def _format_path(path: Path) -> str:
    try:
        return str(path.resolve())
    except Exception:
        return str(path)


def _decode_text(raw: bytes) -> tuple[str, str]:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(encoding), encoding
        except Exception:
            continue
    return raw.decode("utf-8", errors="replace"), "utf-8(replace)"


def _load_file_bytes(path: Path, *, max_bytes: int) -> tuple[bytes, int, bool]:
    total_size = path.stat().st_size
    if total_size <= max_bytes:
        return path.read_bytes(), total_size, False
    with path.open("rb") as handle:
        head = handle.read(max_bytes)
    return head, total_size, True


def _truncate_chars(text: str, max_chars: int) -> tuple[str, bool]:
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars].rstrip(), True


def _query_tokens(query: str) -> tuple[str, ...]:
    return tuple(token.lower() for token in _TOKEN_RE.findall(query) if token)


def _parse_extensions_csv(value: str | None) -> set[str]:
    if not value:
        return set()
    parts = [part.strip().lower() for part in value.split(",") if part.strip()]
    normalized: set[str] = set()
    for part in parts:
        if not part.startswith("."):
            part = "." + part
        normalized.add(part)
    return normalized


def _matches_filename(filename: str, *, query: str, tokens: tuple[str, ...]) -> tuple[bool, int]:
    lowered = filename.lower()
    q = query.lower().strip()
    if not q:
        return False, 0

    score = 0
    if q in lowered:
        score += 4
    matched_tokens = 0
    for token in tokens:
        if token in lowered:
            matched_tokens += 1
    score += matched_tokens

    if matched_tokens == 0 and q not in lowered:
        return False, 0
    return True, score


def _format_file_size(size: int) -> str:
    if size < 1024:
        return f"{size} B"
    if size < 1024 * 1024:
        return f"{size / 1024:.1f} KB"
    return f"{size / (1024 * 1024):.1f} MB"


class _LocalArtifactTool(Tool):
    def _settings(self) -> Any:
        return _config.get()

    def _roots(self) -> tuple[tuple[Path, ...], tuple[Path, ...]]:
        settings = self._settings()
        return _safe_roots(settings), _deny_roots(settings)

    def _max_chars(self, requested: int | None = None) -> int:
        settings = self._settings()
        default = _clamp_int(
            getattr(settings, "local_artifact_max_chars", _DEFAULT_MAX_CHARS),
            _DEFAULT_MAX_CHARS,
            low=500,
            high=_MAX_ALLOWED_CHARS,
        )
        if requested is None:
            return default
        return _clamp_int(requested, default, low=500, high=_MAX_ALLOWED_CHARS)

    def _max_file_bytes(self) -> int:
        settings = self._settings()
        return _clamp_int(
            getattr(settings, "local_artifact_max_file_bytes", _DEFAULT_MAX_FILE_BYTES),
            _DEFAULT_MAX_FILE_BYTES,
            low=1024,
            high=64 * 1024 * 1024,
        )

    def _blocked(self, message: str, *, structured: dict[str, Any] | None = None) -> ToolResult:
        return _tool_result(
            tool_name=self.name,
            content=message,
            is_error=True,
            status="blocked",
            structured=structured,
            retryable=False,
        )

    def _error(
        self,
        message: str,
        *,
        structured: dict[str, Any] | None = None,
        retryable: bool = False,
    ) -> ToolResult:
        return _tool_result(
            tool_name=self.name,
            content=message,
            is_error=True,
            structured=structured,
            retryable=retryable,
        )

    def _allowed_file(self, path: Path) -> tuple[bool, str]:
        safe_roots, deny_roots = self._roots()
        return _allow_path(path, safe_roots=safe_roots, deny_roots=deny_roots)

    def _resolve_file(self, path: str) -> tuple[Path | None, ToolResult | None]:
        candidate = _candidate_path(path)
        if candidate is None:
            return None, self._error("Missing 'path'. Provide a local file path.")
        allowed, reason = self._allowed_file(candidate)
        if not allowed:
            return None, self._blocked(reason, structured={"path": _format_path(candidate)})
        if not candidate.exists():
            return None, self._error(f"File not found: {_format_path(candidate)}")
        if not candidate.is_file():
            return None, self._error(f"Path is not a file: {_format_path(candidate)}")
        return candidate, None


class LocalFindFilesTool(_LocalArtifactTool):
    name = "local_find_files"
    description = "Find local files by name under safe roots (read-only)"
    parameters = {
        "type": "object",
        "properties": {
            "query": {
                "type": "string",
                "description": "Filename query (example: Lasku or invoice pdf)",
            },
            "directory": {
                "type": "string",
                "description": "Optional directory to search within safe roots",
            },
            "max_results": {
                "type": "integer",
                "description": "Max files to return (1-20, default 8)",
            },
            "extensions_csv": {
                "type": "string",
                "description": "Optional comma-separated extensions (example: pdf,docx,xlsx)",
            },
        },
        "required": ["query"],
    }

    def execute(
        self,
        *,
        query: str,
        directory: str | None = None,
        max_results: int = _DEFAULT_FIND_MAX_RESULTS,
        extensions_csv: str | None = None,
        **kwargs: Any,
    ) -> ToolResult:
        del kwargs
        search_query = str(query or "").strip()
        if not search_query:
            return self._error("Empty query. Provide part of the filename to search.")

        safe_roots, deny_roots = self._roots()
        roots: list[Path] = []
        if directory:
            base = _candidate_path(directory)
            if base is None:
                return self._error("Missing 'directory'. Provide a local directory path.")
            allowed, reason = _allow_path(base, safe_roots=safe_roots, deny_roots=deny_roots)
            if not allowed:
                return self._blocked(reason, structured={"directory": _format_path(base)})
            if not base.exists():
                return self._error(f"Directory not found: {_format_path(base)}")
            if not base.is_dir():
                return self._error(f"Path is not a directory: {_format_path(base)}")
            roots = [base]
        else:
            roots = [root for root in safe_roots if root.exists() and root.is_dir()]

        if not roots:
            return self._error("No readable search roots are available.")

        extensions = _parse_extensions_csv(extensions_csv)
        limit = _clamp_int(max_results, _DEFAULT_FIND_MAX_RESULTS, low=1, high=_MAX_FIND_RESULTS)
        max_files_scanned = _clamp_int(
            getattr(self._settings(), "local_artifact_find_max_files", _DEFAULT_FIND_MAX_FILES),
            _DEFAULT_FIND_MAX_FILES,
            low=500,
            high=200000,
        )

        tokens = _query_tokens(search_query)
        matches: list[dict[str, Any]] = []
        scanned = 0
        truncated_scan = False

        for root in roots:
            for dirpath, dirnames, filenames in os.walk(root):
                current = Path(dirpath)

                kept: list[str] = []
                for dirname in dirnames:
                    candidate = current / dirname
                    reason = _path_access_reason(candidate, safe_roots=safe_roots, deny_roots=deny_roots)
                    if reason is None:
                        kept.append(dirname)
                dirnames[:] = kept

                for filename in filenames:
                    scanned += 1
                    if scanned > max_files_scanned:
                        truncated_scan = True
                        break

                    path = current / filename
                    reason = _path_access_reason(path, safe_roots=safe_roots, deny_roots=deny_roots)
                    if reason is not None:
                        continue
                    suffix = path.suffix.lower()
                    if extensions and suffix not in extensions:
                        continue

                    ok, score = _matches_filename(filename, query=search_query, tokens=tokens)
                    if not ok:
                        continue

                    try:
                        stat = path.stat()
                    except OSError:
                        continue

                    matches.append(
                        {
                            "path": _format_path(path),
                            "name": filename,
                            "size_bytes": int(stat.st_size),
                            "mtime": float(stat.st_mtime),
                            "score": score,
                        }
                    )

                if truncated_scan:
                    break
            if truncated_scan:
                break

        matches.sort(key=lambda item: (-int(item.get("score", 0)), -float(item.get("mtime", 0.0))))
        selected = matches[:limit]

        if not selected:
            ext_note = f" with extensions {sorted(extensions)}" if extensions else ""
            return _tool_result(
                tool_name=self.name,
                content=f"No files found for query '{search_query}'{ext_note}.",
                structured={
                    "query": search_query,
                    "directory": _format_path(_candidate_path(directory) or Path(directory)) if directory else None,
                    "scanned_files": scanned,
                    "truncated_scan": truncated_scan,
                    "matches": [],
                },
                sources=[],
                retryable=False,
            )

        lines = [f"Found {len(selected)} file(s) for '{search_query}':"]
        for index, item in enumerate(selected, start=1):
            lines.append(
                f"{index}. {item['path']} ({_format_file_size(int(item['size_bytes']))})"
            )
        if truncated_scan:
            lines.append(
                f"Search stopped early after scanning {max_files_scanned} files. Narrow the directory or query for exhaustive results."
            )

        structured = {
            "query": search_query,
            "directory": _format_path(_candidate_path(directory) or Path(directory)) if directory else None,
            "scanned_files": scanned,
            "truncated_scan": truncated_scan,
            "matches": selected,
        }
        sources = [str(item["path"]) for item in selected]
        return _tool_result(
            tool_name=self.name,
            content="\n".join(lines),
            structured=structured,
            sources=sources,
            retryable=False,
        )


class LocalReadTextTool(_LocalArtifactTool):
    name = "local_read_text"
    description = "Read text from a local text file under safe roots (read-only)"
    parameters = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Path to local text-like file"},
            "max_chars": {
                "type": "integer",
                "description": "Max extracted characters (500-20000)",
            },
        },
        "required": ["path"],
    }

    def execute(self, *, path: str, max_chars: int | None = None, **kwargs: Any) -> ToolResult:
        del kwargs
        resolved, err = self._resolve_file(path)
        if err is not None:
            return err
        assert resolved is not None

        suffix = resolved.suffix.lower()
        if suffix in _PDF_EXTENSIONS:
            return self._error(
                "This file is a PDF. Use local_extract_pdf for reliable extraction.",
                structured={"path": _format_path(resolved), "file_type": "pdf"},
            )
        if suffix in _DOCX_EXTENSIONS:
            return self._error(
                "This file is a DOCX document. Use local_extract_docx for reliable extraction.",
                structured={"path": _format_path(resolved), "file_type": "docx"},
            )
        if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
            return self._error(
                "This file is a spreadsheet workbook. Use local_extract_sheet for reliable extraction.",
                structured={"path": _format_path(resolved), "file_type": "xlsx"},
            )

        limit = self._max_chars(max_chars)
        raw, total_size, byte_truncated = _load_file_bytes(resolved, max_bytes=self._max_file_bytes())
        if b"\x00" in raw[:4096]:
            return self._error(
                "File appears binary and cannot be read as text. Use a format-specific extractor when available.",
                structured={"path": _format_path(resolved), "bytes_read": len(raw), "size_bytes": total_size},
            )

        text, encoding = _decode_text(raw)
        text = text.replace("\r\n", "\n")
        body, char_truncated = _truncate_chars(text, limit)
        truncated = byte_truncated or char_truncated

        structured = {
            "path": _format_path(resolved),
            "file_type": suffix.lstrip(".") or "text",
            "encoding": encoding,
            "size_bytes": total_size,
            "bytes_read": len(raw),
            "chars_returned": len(body),
            "truncated": truncated,
        }
        content = f"Path: {_format_path(resolved)}\n\n{body}"
        if truncated:
            content += "\n\n[Output truncated to safety limits.]"
        return _tool_result(
            tool_name=self.name,
            content=content,
            structured=structured,
            sources=[_format_path(resolved)],
            retryable=False,
        )


class LocalExtractPdfTool(_LocalArtifactTool):
    name = "local_extract_pdf"
    description = "Extract readable text from a local PDF file under safe roots"
    parameters = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Path to local PDF"},
            "max_chars": {
                "type": "integer",
                "description": "Max extracted characters (500-20000)",
            },
            "max_pages": {
                "type": "integer",
                "description": "Max pages to parse (1-120, default 30)",
            },
        },
        "required": ["path"],
    }

    def execute(
        self,
        *,
        path: str,
        max_chars: int | None = None,
        max_pages: int = _DEFAULT_PDF_MAX_PAGES,
        **kwargs: Any,
    ) -> ToolResult:
        del kwargs
        if _PdfReader is None:
            return self._error(
                "PDF extraction dependency missing. Install `pypdf`.",
                structured={"dependency": "pypdf"},
            )

        resolved, err = self._resolve_file(path)
        if err is not None:
            return err
        assert resolved is not None

        if resolved.suffix.lower() not in _PDF_EXTENSIONS:
            return self._error(
                f"File is not a PDF: {_format_path(resolved)}",
                structured={"path": _format_path(resolved), "file_type": resolved.suffix.lower()},
            )

        max_chars_limit = self._max_chars(max_chars)
        page_limit = _clamp_int(max_pages, _DEFAULT_PDF_MAX_PAGES, low=1, high=_MAX_PDF_PAGES)

        try:
            with resolved.open("rb") as handle:
                reader = _PdfReader(handle)
                total_pages = len(reader.pages)
                pages_to_read = min(total_pages, page_limit)
                chunks: list[str] = []
                for idx in range(pages_to_read):
                    page = reader.pages[idx]
                    text = (page.extract_text() or "").strip()
                    if text:
                        chunks.append(f"[Page {idx + 1}]\n{text}")
        except Exception as exc:
            return self._error(
                f"Failed to parse PDF: {exc}",
                structured={"path": _format_path(resolved)},
                retryable=True,
            )

        combined = "\n\n".join(chunks).strip()
        if not combined:
            return self._error(
                "No extractable text found in PDF (it may be scanned images).",
                structured={
                    "path": _format_path(resolved),
                    "pages_read": pages_to_read,
                    "total_pages": total_pages,
                },
            )

        body, char_truncated = _truncate_chars(combined, max_chars_limit)
        structured = {
            "path": _format_path(resolved),
            "file_type": "pdf",
            "total_pages": total_pages,
            "pages_read": pages_to_read,
            "chars_returned": len(body),
            "truncated": bool(char_truncated or pages_to_read < total_pages),
        }
        content = (
            f"Path: {_format_path(resolved)}\n"
            f"Type: pdf\n"
            f"Pages: {pages_to_read}/{total_pages}\n\n"
            f"{body}"
        )
        if char_truncated or pages_to_read < total_pages:
            content += "\n\n[Output truncated to limits.]"
        return _tool_result(
            tool_name=self.name,
            content=content,
            structured=structured,
            sources=[_format_path(resolved)],
            retryable=False,
        )


class LocalExtractDocxTool(_LocalArtifactTool):
    name = "local_extract_docx"
    description = "Extract readable text from a local DOCX file under safe roots"
    parameters = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Path to local DOCX"},
            "max_chars": {
                "type": "integer",
                "description": "Max extracted characters (500-20000)",
            },
        },
        "required": ["path"],
    }

    def execute(self, *, path: str, max_chars: int | None = None, **kwargs: Any) -> ToolResult:
        del kwargs
        if _DocxDocument is None:
            return self._error(
                "DOCX extraction dependency missing. Install `python-docx`.",
                structured={"dependency": "python-docx"},
            )

        resolved, err = self._resolve_file(path)
        if err is not None:
            return err
        assert resolved is not None

        if resolved.suffix.lower() not in _DOCX_EXTENSIONS:
            return self._error(
                f"File is not a DOCX document: {_format_path(resolved)}",
                structured={"path": _format_path(resolved), "file_type": resolved.suffix.lower()},
            )

        try:
            doc = _DocxDocument(_format_path(resolved))
        except Exception as exc:
            return self._error(
                f"Failed to parse DOCX: {exc}",
                structured={"path": _format_path(resolved)},
                retryable=True,
            )

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        table_lines: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if values:
                    table_lines.append(" | ".join(values))

        combined_parts = []
        if paragraphs:
            combined_parts.append("\n".join(paragraphs))
        if table_lines:
            combined_parts.append("\n".join(table_lines))
        combined = "\n\n".join(combined_parts).strip()
        if not combined:
            return self._error(
                "No extractable text found in DOCX.",
                structured={"path": _format_path(resolved)},
            )

        limit = self._max_chars(max_chars)
        body, truncated = _truncate_chars(combined, limit)
        structured = {
            "path": _format_path(resolved),
            "file_type": "docx",
            "paragraph_count": len(paragraphs),
            "table_line_count": len(table_lines),
            "chars_returned": len(body),
            "truncated": truncated,
        }
        content = f"Path: {_format_path(resolved)}\nType: docx\n\n{body}"
        if truncated:
            content += "\n\n[Output truncated to limits.]"

        return _tool_result(
            tool_name=self.name,
            content=content,
            structured=structured,
            sources=[_format_path(resolved)],
            retryable=False,
        )


class LocalExtractSheetTool(_LocalArtifactTool):
    name = "local_extract_sheet"
    description = "Extract readable rows from local CSV/TSV/XLSX files under safe roots"
    parameters = {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Path to CSV/TSV/XLSX file"},
            "max_chars": {
                "type": "integer",
                "description": "Max extracted characters (500-20000)",
            },
            "max_rows": {
                "type": "integer",
                "description": "Max rows per file/sheet (1-500, default 120)",
            },
            "max_sheets": {
                "type": "integer",
                "description": "Max sheets for workbook extraction (1-8, default 3)",
            },
        },
        "required": ["path"],
    }

    def execute(
        self,
        *,
        path: str,
        max_chars: int | None = None,
        max_rows: int = _DEFAULT_SHEET_MAX_ROWS,
        max_sheets: int = _DEFAULT_SHEET_MAX_SHEETS,
        **kwargs: Any,
    ) -> ToolResult:
        del kwargs
        resolved, err = self._resolve_file(path)
        if err is not None:
            return err
        assert resolved is not None

        suffix = resolved.suffix.lower()
        if suffix not in _SHEET_EXTENSIONS:
            return self._error(
                f"Unsupported sheet format for {suffix or 'unknown extension'}. Use csv, tsv, or xlsx files.",
                structured={"path": _format_path(resolved), "file_type": suffix},
            )

        row_limit = _clamp_int(max_rows, _DEFAULT_SHEET_MAX_ROWS, low=1, high=_MAX_SHEET_ROWS)
        sheet_limit = _clamp_int(max_sheets, _DEFAULT_SHEET_MAX_SHEETS, low=1, high=_MAX_SHEET_SHEETS)
        limit = self._max_chars(max_chars)

        try:
            if suffix in {".csv", ".tsv"}:
                lines, details = self._extract_delimited(
                    path=resolved,
                    delimiter="," if suffix == ".csv" else "\t",
                    max_rows=row_limit,
                )
            else:
                lines, details = self._extract_workbook(
                    path=resolved,
                    max_rows=row_limit,
                    max_sheets=sheet_limit,
                )
        except Exception as exc:
            return self._error(
                f"Failed to parse sheet file: {exc}",
                structured={"path": _format_path(resolved), "file_type": suffix},
                retryable=True,
            )

        if not lines:
            return self._error(
                "No extractable rows found in sheet.",
                structured={"path": _format_path(resolved), "file_type": suffix},
            )

        combined = "\n".join(lines)
        body, truncated = _truncate_chars(combined, limit)
        structured = {
            "path": _format_path(resolved),
            "file_type": suffix.lstrip("."),
            "chars_returned": len(body),
            "truncated": truncated,
            **details,
        }
        content = f"Path: {_format_path(resolved)}\nType: {suffix.lstrip('.')}\n\n{body}"
        if truncated:
            content += "\n\n[Output truncated to limits.]"
        return _tool_result(
            tool_name=self.name,
            content=content,
            structured=structured,
            sources=[_format_path(resolved)],
            retryable=False,
        )

    def _extract_delimited(self, *, path: Path, delimiter: str, max_rows: int) -> tuple[list[str], dict[str, Any]]:
        raw, total_size, byte_truncated = _load_file_bytes(path, max_bytes=self._max_file_bytes())
        text, encoding = _decode_text(raw)
        stream = io.StringIO(text)
        reader = csv.reader(stream, delimiter=delimiter)

        lines: list[str] = []
        count = 0
        for row in reader:
            if count >= max_rows:
                break
            cleaned = [_normalize_space(str(cell)) for cell in row]
            if any(cleaned):
                lines.append(" | ".join(cleaned))
            count += 1

        details = {
            "encoding": encoding,
            "rows_read": count,
            "rows_limit": max_rows,
            "size_bytes": total_size,
            "bytes_truncated": byte_truncated,
            "sheet_names": [path.stem],
        }
        return lines, details

    def _extract_workbook(self, *, path: Path, max_rows: int, max_sheets: int) -> tuple[list[str], dict[str, Any]]:
        if _load_workbook is None:
            raise RuntimeError("XLSX extraction dependency missing. Install `openpyxl`.")

        workbook = _load_workbook(filename=_format_path(path), read_only=True, data_only=True)
        lines: list[str] = []
        sheet_names: list[str] = []
        rows_read = 0
        sheets_read = 0

        try:
            for worksheet in workbook.worksheets:
                if sheets_read >= max_sheets:
                    break
                sheets_read += 1
                sheet_names.append(str(worksheet.title))
                lines.append(f"[Sheet: {worksheet.title}]")

                row_count = 0
                for row in worksheet.iter_rows(values_only=True):
                    if row_count >= max_rows:
                        break
                    cleaned = [_normalize_space(str(cell)) for cell in row if cell is not None]
                    if cleaned:
                        lines.append(" | ".join(cleaned))
                    row_count += 1
                    rows_read += 1
                lines.append("")
        finally:
            workbook.close()

        details = {
            "sheet_names": sheet_names,
            "sheets_read": sheets_read,
            "sheets_limit": max_sheets,
            "rows_read": rows_read,
            "rows_limit": max_rows,
        }
        return lines, details
